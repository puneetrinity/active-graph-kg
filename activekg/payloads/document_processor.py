"""
RAG Document Processing System
Implements document processing, chunking, and storage for RAG workflows
"""

import json
import sqlite3
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, cast

import numpy as np

from activekg.common.logger import get_enhanced_logger

logger = get_enhanced_logger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of a document with metadata and embeddings"""

    chunk_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    content: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)
    embedding: np.ndarray | None = None
    source_document_id: str = ""
    chunk_index: int = 0
    chunk_type: str = "text"  # text, table, image, code, etc.
    relevance_score: float = 0.0
    created_at: datetime = field(default_factory=datetime.utcnow)

    def to_dict(self) -> dict[str, Any]:
        """Convert chunk to dictionary for serialization"""
        return {
            "chunk_id": self.chunk_id,
            "content": self.content,
            "metadata": self.metadata,
            "source_document_id": self.source_document_id,
            "chunk_index": self.chunk_index,
            "chunk_type": self.chunk_type,
            "relevance_score": self.relevance_score,
            "created_at": self.created_at.isoformat(),
            "embedding_shape": self.embedding.shape if self.embedding is not None else None,
        }

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "DocumentChunk":
        """Create chunk from dictionary"""
        chunk = cls(
            chunk_id=data["chunk_id"],
            content=data["content"],
            metadata=data.get("metadata", {}),
            source_document_id=data["source_document_id"],
            chunk_index=data["chunk_index"],
            chunk_type=data.get("chunk_type", "text"),
            relevance_score=data.get("relevance_score", 0.0),
            created_at=datetime.fromisoformat(data["created_at"]),
        )
        return chunk


@dataclass
class Document:
    """Represents a document with metadata and chunks"""

    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    filename: str = ""
    content: str = ""
    content_type: str = ""
    file_size: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)
    upload_date: datetime = field(default_factory=datetime.utcnow)
    processed_date: datetime | None = None
    chunks: list[DocumentChunk] = field(default_factory=list)
    status: str = "pending"  # pending, processing, completed, error

    def to_dict(self) -> dict[str, Any]:
        """Convert document to dictionary for serialization"""
        return {
            "id": self.id,
            "filename": self.filename,
            "content_type": self.content_type,
            "file_size": self.file_size,
            "metadata": self.metadata,
            "upload_date": self.upload_date.isoformat(),
            "processed_date": self.processed_date.isoformat() if self.processed_date else None,
            "chunk_count": len(self.chunks),
            "status": self.status,
        }


class DocumentProcessor:
    """Handles document processing for various file types"""

    def __init__(self):
        self.supported_types = {
            ".pdf",
            ".txt",
            ".docx",
            ".doc",
            ".html",
            ".md",
            ".json",
            ".xlsx",
            ".xls",
            ".csv",
        }
        self.logger = logger

    def process_document(self, content: str | bytes, filename: str, content_type: str) -> Document:
        """
        Process a document and extract text content

        Args:
            content: Document content as string or bytes
            filename: Original filename
            content_type: MIME type or file extension

        Returns:
            Document object with extracted content
        """
        self.logger.info(f"Processing document: {filename}")

        try:
            # Extract text based on content type
            if content_type.lower() in [".txt", ".md", "text/plain"]:
                text_content = self._process_text(content)
            elif content_type.lower() == ".json":
                text_content = self._process_json(content)
            elif content_type.lower() == ".html":
                text_content = self._process_html(content)
            elif content_type.lower() == ".pdf":
                text_content = self._process_pdf(content)
            elif content_type.lower() in [".docx", ".doc"]:
                text_content = self._process_docx(content)
            elif content_type.lower() in [".xlsx", ".xls"]:
                text_content = self._process_excel(content)
            elif content_type.lower() == ".csv":
                text_content = self._process_csv(content)
            else:
                # Default to treating as text
                text_content = self._process_text(content)

            # Create document object
            document = Document(
                filename=filename,
                content=text_content,
                content_type=content_type,
                file_size=len(content) if isinstance(content, (str, bytes)) else 0,
                processed_date=datetime.utcnow(),
                status="processing",
            )

            self.logger.info(f"Document processed successfully: {filename}")
            return document

        except Exception as e:
            self.logger.error(f"Error processing document {filename}: {e}")
            raise

    def _process_text(self, content: str | bytes) -> str:
        """Process plain text content"""
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="ignore")
        return self._clean_text(content)

    def _process_json(self, content: str | bytes) -> str:
        """Process JSON content by extracting text values"""
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="ignore")

        try:
            data = json.loads(content)
            text_parts = []

            def extract_text_from_json(obj):
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, str):
                            text_parts.append(f"{key}: {value}")
                        elif isinstance(value, (dict, list)):
                            extract_text_from_json(value)
                elif isinstance(obj, list):
                    for item in obj:
                        extract_text_from_json(item)
                elif isinstance(obj, str):
                    text_parts.append(obj)

            extract_text_from_json(data)
            return self._clean_text("\n".join(text_parts))

        except json.JSONDecodeError:
            return self._process_text(content)

    def _process_html(self, content: str | bytes) -> str:
        """Process HTML content by extracting text"""
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="ignore")

        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(content, "html.parser")
            text = soup.get_text(separator=" ", strip=True)
            return self._clean_text(text)
        except ImportError:
            # Fallback if BeautifulSoup is not available
            import re

            text = re.sub(r"<[^>]+>", " ", content)
            return self._clean_text(text)

    def _process_pdf(self, content: str | bytes) -> str:
        """Process PDF content and extract text"""
        if isinstance(content, str):
            content = content.encode("utf-8")

        try:
            import io

            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                full_text = "\n".join(text_parts)
                return self._clean_text(full_text)

        except ImportError:
            # Fallback to PyPDF2
            try:
                import io

                import PyPDF2

                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_parts = []

                for page in pdf_reader.pages:
                    # PyPDF2 type stubs have inconsistencies, cast to Any to avoid false positives
                    page_text = cast(Any, page).extract_text()
                    if page_text:
                        text_parts.append(page_text)

                full_text = "\n".join(text_parts)
                return self._clean_text(full_text)

            except Exception as e:
                self.logger.error(f"Error processing PDF: {e}")
                return "Error: Could not extract text from PDF"

        except Exception as e:
            self.logger.error(f"Error processing PDF with pdfplumber: {e}")
            return "Error: Could not extract text from PDF"

    def _process_docx(self, content: str | bytes) -> str:
        """Process DOCX/DOC content and extract text"""
        if isinstance(content, str):
            content = content.encode("utf-8")

        try:
            import io

            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(content))
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            full_text = "\n".join(text_parts)
            return self._clean_text(full_text)

        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            return "Error: Could not extract text from DOCX"

    def _process_excel(self, content: str | bytes) -> str:
        """Process Excel content and extract text"""
        if isinstance(content, str):
            content = content.encode("utf-8")

        try:
            import io

            import openpyxl

            workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(row_values).strip()
                    if row_text and row_text != " | " * (len(row_values) - 1):
                        text_parts.append(row_text)

            full_text = "\n".join(text_parts)
            return self._clean_text(full_text)

        except Exception as e:
            # Fallback to xlrd for older Excel formats
            try:
                import io

                import xlrd

                workbook = xlrd.open_workbook(file_contents=content)
                text_parts = []

                for sheet_idx in range(workbook.nsheets):
                    sheet = workbook.sheet_by_index(sheet_idx)
                    text_parts.append(f"Sheet: {sheet.name}")

                    for row_idx in range(sheet.nrows):
                        row_values = [
                            str(sheet.cell_value(row_idx, col_idx))
                            for col_idx in range(sheet.ncols)
                        ]
                        row_text = " | ".join(row_values).strip()
                        if row_text:
                            text_parts.append(row_text)

                full_text = "\n".join(text_parts)
                return self._clean_text(full_text)

            except Exception as e2:
                self.logger.error(f"Error processing Excel: {e}, {e2}")
                return "Error: Could not extract text from Excel file"

    def _process_csv(self, content: str | bytes) -> str:
        """Process CSV content and extract text"""
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="ignore")

        try:
            import csv
            import io

            csv_reader = csv.reader(io.StringIO(content))
            text_parts = []

            for row in csv_reader:
                row_text = " | ".join(str(cell).strip() for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)

            full_text = "\n".join(text_parts)
            return self._clean_text(full_text)

        except Exception as e:
            self.logger.error(f"Error processing CSV: {e}")
            return self._process_text(content)  # Fallback to text processing

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = " ".join(text.split())

        # Remove special characters that might interfere
        text = text.replace("\x00", "")  # Remove null bytes
        text = text.replace("\ufeff", "")  # Remove BOM

        # Remove very long lines that might be corrupted
        lines = text.split("\n")
        cleaned_lines = [line for line in lines if len(line) < 10000]

        return "\n".join(cleaned_lines).strip()


class DocumentChunker:
    """Handles document chunking with various strategies"""

    def __init__(self, chunk_size: int = 512, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.logger = logger

    def chunk_document(self, document: Document, strategy: str = "semantic") -> list[DocumentChunk]:
        """
        Split document into chunks using specified strategy

        Args:
            document: Document to chunk
            strategy: Chunking strategy ('semantic', 'fixed', 'paragraph')

        Returns:
            List of DocumentChunk objects
        """
        self.logger.info(f"Chunking document {document.filename} with strategy: {strategy}")

        try:
            if strategy == "semantic":
                chunks = self._semantic_chunk(document)
            elif strategy == "fixed":
                chunks = self._fixed_size_chunk(document)
            elif strategy == "paragraph":
                chunks = self._paragraph_chunk(document)
            else:
                self.logger.warning(f"Unknown chunking strategy: {strategy}, using semantic")
                chunks = self._semantic_chunk(document)

            # Set chunk metadata
            for i, chunk in enumerate(chunks):
                chunk.source_document_id = document.id
                chunk.chunk_index = i
                chunk.metadata = {
                    "source_filename": document.filename,
                    "source_content_type": document.content_type,
                    "chunking_strategy": strategy,
                    "chunk_size": self.chunk_size,
                    "overlap": self.overlap,
                }

            self.logger.info(f"Created {len(chunks)} chunks for document {document.filename}")
            return chunks

        except Exception as e:
            self.logger.error(f"Error chunking document {document.filename}: {e}")
            raise

    def _semantic_chunk(self, document: Document) -> list[DocumentChunk]:
        """Split text into semantically coherent chunks"""
        text = document.content
        sentences = self._split_into_sentences(text)
        chunks = []
        current_chunk = ""
        current_size = 0

        for sentence in sentences:
            sentence_len = len(sentence)

            # If adding this sentence would exceed chunk size, create a new chunk
            if current_size + sentence_len > self.chunk_size and current_chunk:
                chunks.append(DocumentChunk(content=current_chunk.strip()))

                # Handle overlap
                if self.overlap > 0:
                    current_chunk = current_chunk[-self.overlap :] + " " + sentence
                    current_size = len(current_chunk)
                else:
                    current_chunk = sentence
                    current_size = sentence_len
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_size += sentence_len

        # Add final chunk if there's content
        if current_chunk.strip():
            chunks.append(DocumentChunk(content=current_chunk.strip()))

        return chunks

    def _fixed_size_chunk(self, document: Document) -> list[DocumentChunk]:
        """Split text into fixed-size chunks with overlap"""
        text = document.content
        chunks = []

        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunk_text = text[start:end]

            if chunk_text.strip():
                chunks.append(DocumentChunk(content=chunk_text.strip()))

            start = end - self.overlap if self.overlap > 0 else end

        return chunks

    def _paragraph_chunk(self, document: Document) -> list[DocumentChunk]:
        """Split text into chunks by paragraphs"""
        text = document.content
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # If adding this paragraph would exceed chunk size, create a new chunk
            if len(current_chunk) + len(paragraph) > self.chunk_size and current_chunk:
                chunks.append(DocumentChunk(content=current_chunk.strip()))
                current_chunk = paragraph
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph

        # Add final chunk if there's content
        if current_chunk.strip():
            chunks.append(DocumentChunk(content=current_chunk.strip()))

        return chunks

    def _split_into_sentences(self, text: str) -> list[str]:
        """Split text into sentences using simple heuristics"""
        import re

        # Simple sentence splitting using regex
        sentences = re.split(r"[.!?]+", text)

        # Filter out empty sentences and clean up
        sentences = [s.strip() for s in sentences if s.strip()]

        return sentences


class DocumentStore:
    """Handles document storage and retrieval"""

    def __init__(
        self, db_path: str = "data/rag_documents.db", documents_dir: str = "data/documents"
    ):
        self.db_path = db_path
        self.documents_dir = Path(documents_dir)
        self.documents_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logger
        self._init_database()

    def _init_database(self):
        """Initialize SQLite database for document metadata"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute("""
                    CREATE TABLE IF NOT EXISTS documents (
                        id TEXT PRIMARY KEY,
                        filename TEXT NOT NULL,
                        content_type TEXT,
                        file_size INTEGER,
                        upload_date TEXT,
                        processed_date TEXT,
                        metadata TEXT,
                        chunk_count INTEGER DEFAULT 0,
                        status TEXT DEFAULT 'pending'
                    )
                """)

                conn.execute("""
                    CREATE TABLE IF NOT EXISTS document_chunks (
                        chunk_id TEXT PRIMARY KEY,
                        document_id TEXT,
                        chunk_index INTEGER,
                        content TEXT,
                        metadata TEXT,
                        relevance_score REAL DEFAULT 0.0,
                        created_at TEXT,
                        FOREIGN KEY (document_id) REFERENCES documents (id)
                    )
                """)

                # Create indexes for performance
                conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_chunks_document_id
                    ON document_chunks(document_id)
                """)

                conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_chunks_relevance
                    ON document_chunks(relevance_score DESC)
                """)

                conn.commit()

        except Exception as e:
            self.logger.error(f"Error initializing database: {e}")
            raise

    def store_document(self, document: Document, chunks: list[DocumentChunk]) -> bool:
        """Store document and its chunks"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                # Store document metadata
                conn.execute(
                    """
                    INSERT OR REPLACE INTO documents
                    (id, filename, content_type, file_size, upload_date,
                     processed_date, metadata, chunk_count, status)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                    (
                        document.id,
                        document.filename,
                        document.content_type,
                        document.file_size,
                        document.upload_date.isoformat(),
                        document.processed_date.isoformat() if document.processed_date else None,
                        json.dumps(document.metadata),
                        len(chunks),
                        document.status,
                    ),
                )

                # Store chunks
                for chunk in chunks:
                    conn.execute(
                        """
                        INSERT OR REPLACE INTO document_chunks
                        (chunk_id, document_id, chunk_index, content, metadata, created_at)
                        VALUES (?, ?, ?, ?, ?, ?)
                    """,
                        (
                            chunk.chunk_id,
                            chunk.source_document_id,
                            chunk.chunk_index,
                            chunk.content,
                            json.dumps(chunk.metadata),
                            chunk.created_at.isoformat(),
                        ),
                    )

                # Store full document content separately
                doc_file_path = self.documents_dir / f"{document.id}.json"
                with open(doc_file_path, "w", encoding="utf-8") as f:
                    json.dump(
                        {
                            "document": document.to_dict(),
                            "content": document.content,
                            "chunks": [chunk.to_dict() for chunk in chunks],
                        },
                        f,
                        ensure_ascii=False,
                        indent=2,
                    )

                conn.commit()

            self.logger.info(f"Stored document {document.id} with {len(chunks)} chunks")
            return True

        except Exception as e:
            self.logger.error(f"Error storing document: {e}")
            return False

    def retrieve_document(self, document_id: str) -> Document | None:
        """Retrieve document by ID"""
        try:
            doc_file_path = self.documents_dir / f"{document_id}.json"
            if not doc_file_path.exists():
                return None

            with open(doc_file_path, encoding="utf-8") as f:
                data = json.load(f)

            doc_data = data["document"]
            document = Document(
                id=doc_data["id"],
                filename=doc_data["filename"],
                content=data["content"],
                content_type=doc_data["content_type"],
                file_size=doc_data["file_size"],
                metadata=doc_data["metadata"],
                upload_date=datetime.fromisoformat(doc_data["upload_date"]),
                processed_date=datetime.fromisoformat(doc_data["processed_date"])
                if doc_data["processed_date"]
                else None,
                status=doc_data["status"],
            )

            # Load chunks
            for chunk_data in data["chunks"]:
                chunk = DocumentChunk.from_dict(chunk_data)
                document.chunks.append(chunk)

            return document

        except Exception as e:
            self.logger.error(f"Error retrieving document {document_id}: {e}")
            return None

    def get_chunks_by_document_id(self, document_id: str) -> list[DocumentChunk]:
        """Get all chunks for a document"""
        chunks = []
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    """
                    SELECT chunk_id, chunk_index, content, metadata, created_at
                    FROM document_chunks
                    WHERE document_id = ?
                    ORDER BY chunk_index
                """,
                    (document_id,),
                )

                for row in cursor.fetchall():
                    chunk = DocumentChunk(
                        chunk_id=row[0],
                        source_document_id=document_id,
                        chunk_index=row[1],
                        content=row[2],
                        metadata=json.loads(row[3]),
                        created_at=datetime.fromisoformat(row[4]),
                    )
                    chunks.append(chunk)

        except Exception as e:
            self.logger.error(f"Error retrieving chunks for document {document_id}: {e}")

        return chunks

    def search_documents(self, query: str, limit: int = 10) -> list[dict]:
        """Simple text search across documents"""
        results = []
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    """
                    SELECT id, filename, content_type, upload_date, chunk_count, status
                    FROM documents
                    WHERE filename LIKE ? OR metadata LIKE ?
                    ORDER BY upload_date DESC
                    LIMIT ?
                """,
                    (f"%{query}%", f"%{query}%", limit),
                )

                for row in cursor.fetchall():
                    results.append(
                        {
                            "id": row[0],
                            "filename": row[1],
                            "content_type": row[2],
                            "upload_date": row[3],
                            "chunk_count": row[4],
                            "status": row[5],
                        }
                    )

        except Exception as e:
            self.logger.error(f"Error searching documents: {e}")

        return results

    def list_documents(self, limit: int = 50, offset: int = 0) -> list[dict]:
        """List all documents"""
        results = []
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    """
                    SELECT id, filename, content_type, upload_date, chunk_count, status
                    FROM documents
                    ORDER BY upload_date DESC
                    LIMIT ? OFFSET ?
                """,
                    (limit, offset),
                )

                for row in cursor.fetchall():
                    results.append(
                        {
                            "id": row[0],
                            "filename": row[1],
                            "content_type": row[2],
                            "upload_date": row[3],
                            "chunk_count": row[4],
                            "status": row[5],
                        }
                    )

        except Exception as e:
            self.logger.error(f"Error listing documents: {e}")

        return results

    def delete_document(self, document_id: str) -> bool:
        """Delete a document and its chunks"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                # Delete chunks first
                conn.execute("DELETE FROM document_chunks WHERE document_id = ?", (document_id,))

                # Delete document
                conn.execute("DELETE FROM documents WHERE id = ?", (document_id,))

                # Delete document file
                doc_file_path = self.documents_dir / f"{document_id}.json"
                if doc_file_path.exists():
                    doc_file_path.unlink()

                conn.commit()

            self.logger.info(f"Deleted document {document_id}")
            return True

        except Exception as e:
            self.logger.error(f"Error deleting document {document_id}: {e}")
            return False
